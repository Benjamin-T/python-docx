# encoding: utf-8

"""Test suite for the docx.blkcntnr (block item container) module"""

from __future__ import absolute_import, division, print_function, unicode_literals

import pytest

from docx.blkcntnr import BlockItemContainer
from docx.bookmark import Bookmarks, _DocumentBookmarkFinder
from docx.shared import Inches
from docx.table import Table
from docx.text.paragraph import Paragraph

from .unitutil.cxml import element, xml
from .unitutil.file import snippet_seq
from .unitutil.mock import call, instance_mock, method_mock, property_mock


class DescribeBlockItemContainer(object):
    def it_can_add_a_named_bookmark(self, start_bookmark_fixture):
        next_id_, element_, expected_xml = start_bookmark_fixture

        blkcntnr = BlockItemContainer(element_, None)
        bookmark = blkcntnr.start_bookmark("bmk-1")

        assert blkcntnr._element.xml == expected_xml
        assert bookmark.name == "bmk-1"
        assert bookmark.id == 0
        assert bookmark._bookmarkEnd is None
        next_id_.assert_called_once_with()

    def but_it_raises_KeyError_when_name_exists(self, part, next_id_, bookmark_names_):
        blkcntnr = BlockItemContainer(element("w:body"), None)
        next_id_.return_value = 1
        bookmark_names_.return_value = ["bmk-0", "bmk-2"]

        with pytest.raises(KeyError) as exc:
            blkcntnr.start_bookmark("bmk-0")
        assert "Bookmark name already present in document." in str(exc.value)

    # fixtures -------------------------------------------------------

    @pytest.fixture(
        params=[
            ("w:p", "w:p/(w:bookmarkStart{w:name=bmk-1, w:id=0})"),
            ("w:body", "w:body/(w:bookmarkStart{w:name=bmk-1, w:id=0})"),
            ("w:ftr", "w:ftr/(w:bookmarkStart{w:name=bmk-1, w:id=0})"),
            ("w:hdr", "w:hdr/(w:bookmarkStart{w:name=bmk-1, w:id=0})"),
        ]
    )
    def start_bookmark_fixture(self, request, next_id_, bookmark_names_, part):
        cxml, expected_cxml = request.param
        expected_xml = xml(expected_cxml)
        element_ = element(cxml)
        next_id_.return_value = 0
        bookmark_names_.return_value = ["bmk-0", "bmk-2"]
        return next_id_, element_, expected_xml

    def it_can_add_a_paragraph(self, add_paragraph_fixture, _add_paragraph_):
        text, style, paragraph_, add_run_calls = add_paragraph_fixture
        _add_paragraph_.return_value = paragraph_
        blkcntnr = BlockItemContainer(None, None)

        paragraph = blkcntnr.add_paragraph(text, style)

        _add_paragraph_.assert_called_once_with(blkcntnr)
        assert paragraph.add_run.call_args_list == add_run_calls
        assert paragraph.style == style
        assert paragraph is paragraph_

    def it_can_add_a_table(self, add_table_fixture):
        blkcntnr, rows, cols, width, expected_xml = add_table_fixture
        table = blkcntnr.add_table(rows, cols, width)
        assert isinstance(table, Table)
        assert table._element.xml == expected_xml
        assert table._parent is blkcntnr

    def it_provides_access_to_the_paragraphs_it_contains(self, paragraphs_fixture):
        # test len(), iterable, and indexed access
        blkcntnr, expected_count = paragraphs_fixture
        paragraphs = blkcntnr.paragraphs
        assert len(paragraphs) == expected_count
        count = 0
        for idx, paragraph in enumerate(paragraphs):
            assert isinstance(paragraph, Paragraph)
            assert paragraphs[idx] is paragraph
            count += 1
        assert count == expected_count

    def it_provides_access_to_the_tables_it_contains(self, tables_fixture):
        # test len(), iterable, and indexed access
        blkcntnr, expected_count = tables_fixture
        tables = blkcntnr.tables
        assert len(tables) == expected_count
        count = 0
        for idx, table in enumerate(tables):
            assert isinstance(table, Table)
            assert tables[idx] is table
            count += 1
        assert count == expected_count

    def it_adds_a_paragraph_to_help(self, _add_paragraph_fixture):
        blkcntnr, expected_xml = _add_paragraph_fixture
        new_paragraph = blkcntnr._add_paragraph()
        assert isinstance(new_paragraph, Paragraph)
        assert new_paragraph._parent == blkcntnr
        assert blkcntnr._element.xml == expected_xml

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[("", None), ("Foo", None), ("", "Bar"), ("Foo", "Bar")])
    def add_paragraph_fixture(self, request, paragraph_):
        text, style = request.param
        paragraph_.style = None
        add_run_calls = [call(text)] if text else []
        return text, style, paragraph_, add_run_calls

    @pytest.fixture
    def _add_paragraph_fixture(self, request):
        blkcntnr_cxml, after_cxml = "w:body", "w:body/w:p"
        blkcntnr = BlockItemContainer(element(blkcntnr_cxml), None)
        expected_xml = xml(after_cxml)
        return blkcntnr, expected_xml

    @pytest.fixture
    def add_table_fixture(self):
        blkcntnr = BlockItemContainer(element("w:body"), None)
        rows, cols, width = 2, 2, Inches(2)
        expected_xml = snippet_seq("new-tbl")[0]
        return blkcntnr, rows, cols, width, expected_xml

    @pytest.fixture(
        params=[
            ("w:body", 0),
            ("w:body/w:p", 1),
            ("w:body/(w:p,w:p)", 2),
            ("w:body/(w:p,w:tbl)", 1),
            ("w:body/(w:p,w:tbl,w:p)", 2),
        ]
    )
    def paragraphs_fixture(self, request):
        blkcntnr_cxml, expected_count = request.param
        blkcntnr = BlockItemContainer(element(blkcntnr_cxml), None)
        return blkcntnr, expected_count

    @pytest.fixture(
        params=[
            ("w:body", 0),
            ("w:body/w:tbl", 1),
            ("w:body/(w:tbl,w:tbl)", 2),
            ("w:body/(w:p,w:tbl)", 1),
            ("w:body/(w:tbl,w:tbl,w:p)", 2),
        ]
    )
    def tables_fixture(self, request):
        blkcntnr_cxml, expected_count = request.param
        blkcntnr = BlockItemContainer(element(blkcntnr_cxml), None)
        return blkcntnr, expected_count

    # fixture components ---------------------------------------------

    @pytest.fixture
    def _add_paragraph_(self, request):
        return method_mock(request, BlockItemContainer, "_add_paragraph")

    @pytest.fixture
    def Bookmarks_(self, request):
        return instance_mock(request, Bookmarks)

    @pytest.fixture
    def paragraph_(self, request):
        return instance_mock(request, Paragraph)

    @pytest.fixture
    def part(self, request):
        return property_mock(request, BlockItemContainer, "part")

    @pytest.fixture
    def _DocumentBookmarkFinder_(self, request):
        return instance_mock(request, _DocumentBookmarkFinder)

    @pytest.fixture
    def bookmark_names_(self, request):
        return property_mock(request, _DocumentBookmarkFinder, "bookmark_names")

    @pytest.fixture
    def next_id_(self, request):
        return property_mock(request, _DocumentBookmarkFinder, "next_id")
